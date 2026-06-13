"""Build the UP political-history note as a DOCX per the assessment's formatting rules.

Formatting spec (from the guidelines):
- Font: Arial, size 11; single line spacing; A4; 0.5" margins on all sides.
- One line gap between paragraphs; justified text.
- Header (right): "816258_Divyaman Pal" in Arial size 9.
- No document title inside the note; section headings are allowed.
"""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

OUT = "up_note/Divyaman_Pal_Note writing assessment.docx"
HEADER_TEXT = "816258_Divyaman Pal"
BODY_FONT = "Arial"
BODY_SIZE = 11
GAP_AFTER = Pt(10)  # ~ one line gap between paragraphs at 11pt single spacing

# (heading, [paragraphs]) ; heading None => intro with no heading
CONTENT = [
    (None, [
        "Uttar Pradesh's 403 Assembly and 80 Lok Sabha seats make it the pivot of national power, and its "
        "politics since 1989 form a single storyline: the collapse of the Congress umbrella, its replacement by "
        "two caste-bloc projects \u2013 Mandal (Samajwadi Party) and Bahujan (Bahujan Samaj Party) \u2013 and the "
        "BJP's success after 2014 in fusing non-Yadav OBCs, non-Jatav Dalits and upper castes into one Hindu "
        "coalition powered by Hindutva, welfare delivery and centralised governance. That formula peaked in 2017 "
        "(312/403) and held in 2022 (255), before social-justice arithmetic and governance discontent cut the BJP "
        "to 33 of 80 seats in 2024. The recommendations below follow this arc.",
    ]),
    ("Demographic and Geographic Analysis", [
        "Recommendation: UP is a \u201c20-20-20-40\u201d state \u2013 roughly 20% upper caste, ~21% Dalit, ~19\u201320% "
        "Muslim and 40\u201345% OBC \u2013 so elections are won by whoever assembles the largest cross-caste bloc, "
        "not by any single community. The SP's core is Yadav (9\u201310%) plus Muslim; the BSP's is Jatav (10\u201312%); "
        "both are ~30% ceilings that cannot win alone. The decisive swing groups are the non-Yadav OBCs (~30%: "
        "Kurmi, Maurya/Kushwaha, Lodh, Nishad, Rajbhar) and non-Jatav Dalits (8\u201310%: Pasi, Valmiki).",
        "The evidence tracks this arithmetic precisely. The BJP's dominance rested on consolidating these swing "
        "blocs with upper castes (Brahmin 9\u201310%, Thakur 7\u20138%) behind Hindu unity: 312 seats (39.7%) in "
        "2017 and 255 (41.3%) in 2022, against rivals stuck near their caste ceilings (BSP 22.2% and SP 21.8% in "
        "2017). The 2024 reversal came when the SP's social broadening prised non-Yadav OBCs and non-Jatav Dalits "
        "partly away, lifting the INDIA bloc to 43/80 as the BSP's Dalit vote collapsed to 9.4%. Across 2012\u20132024 "
        "the lesson is constant: caste arithmetic decides UP, and the marginal OBC and Dalit sub-castes are the prize.",
        "Geography sharpens the arithmetic. Western UP and the Terai (sugarcane, dairy, NCR proximity) are the most "
        "prosperous zones, with Jat/RLD farmer politics where MSP, cane arrears and the 2020\u201321 farm-law "
        "agitation move votes; the RLD's alignment (with the SP in 2022, the BJP in 2024) repeatedly tilts a dozen-odd "
        "western seats. The Awadh-central and Purvanchal (eastern) Gangetic plains are the most populous, "
        "caste-saturated battlegrounds that decide the overall verdict. Bundelkhand (drought, debt, stray-cattle "
        "distress, out-migration) and the Vindhyachal/south-east (forest, tribal, mineral, backward) are resource-poor "
        "zones where welfare and basic infrastructure outweigh identity. The practical implication is that parties must "
        "localise: development and expressway messaging in the west, welfare and social-justice messaging in the east "
        "and Bundelkhand.",
    ]),
    ("Recent Administrative Evolution", [
        "Recommendation: district and division reorganisation in UP has been used less for administrative efficiency "
        "than as symbolic identity-signalling and patronage, and it should instead be judged on service-delivery "
        "outcomes. The state now has 75 districts and 18 divisions.",
        "The pattern is bipartisan but caste-coded. Mayawati's BSP carved out about ten districts named after Bahujan "
        "icons \u2013 Sant Ravidas Nagar (1994), Ambedkar Nagar (1995), Mahamaya Nagar (2002), Kanshiram Nagar (2008), "
        "Chhatrapati Shahuji Maharaj Nagar (2010) and, weeks before the 2012 polls, Prabuddh Nagar, Panchsheel Nagar "
        "and Bhim Nagar \u2013 embedding Dalit memory in the map while distributing new district headquarters, posts and "
        "contracts as patronage; she simultaneously demanded trifurcation into Pashchimanchal, Purvanchal and "
        "Bundelkhand. In 2012 the Akhilesh-led SP reversed eight of these names to local toponyms (Kanshiram "
        "Nagar\u2192Kasganj, Mahamaya Nagar\u2192Hathras, Jyotiba Phule Nagar\u2192Amroha, CSM Nagar\u2192Amethi, Rama "
        "Bai Nagar\u2192Kanpur Dehat, Bhim Nagar\u2192Sambhal, Prabuddh Nagar\u2192Shamli, Panchsheel Nagar\u2192Hapur), "
        "which Mayawati attacked as an \u201cinsult to Dalit icons\u201d \u2013 a counter-mobilisation in itself. Since "
        "2017 the BJP has continued the logic through Hindutva-coded renaming \u2013 Allahabad\u2192Prayagraj and "
        "Faizabad\u2192Ayodhya (2018), and Mughalsarai junction\u2192Deen Dayal Upadhyaya Nagar \u2013 converting "
        "administrative cartography into cultural assertion. The governance lesson is that map-making is a low-cost "
        "substitute for delivery and should be replaced by measurable improvements in administration.",
    ]),
    ("Leadership Impact", [
        "Recommendation: Mulayam Singh Yadav and Kanshi Ram permanently democratised UP politics by institutionalising "
        "backward-caste and Dalit assertion, while Yogi Adityanath has re-centralised it around governance-plus-Hindutva; "
        "any opposition strategy must reckon with both legacies.",
        "Mulayam, a Lohiaite socialist who founded the SP in 1992, built the Yadav-Muslim (MY) coalition into a governing "
        "force (Chief Minister 1989\u201391, 1993\u201395, 2003\u201307) and, through the 1993 SP-BSP tie-up, blocked the "
        "BJP immediately after Babri \u2013 making OBC empowerment the central axis of state politics. Kanshi Ram built "
        "the organisational scaffolding of Dalit power \u2013 BAMCEF (1971), DS-4 and the BSP (1984) \u2013 converting "
        "Ambedkarite identity into disciplined electoral strength and mentoring Mayawati, whose 2007 \u201csarvajan\u201d "
        "(Dalit plus Brahmin) experiment delivered a rare single-party majority (206/403). Together they replaced "
        "Congress's patron-client model with competitive identity mobilisation that still structures the state.",
        "Yogi Adityanath, Chief Minister since 2017 and the first UP leader in 37 years to return after a full term (2022), "
        "recast that culture. He pairs a law-and-order narrative (encounters, \u201cbulldozer\u201d action, anti-Romeo "
        "squads) with visible development (the Purvanchal, Bundelkhand and Ganga expressways, new airports, Kashi "
        "Vishwanath and the Ram temple), targeted welfare, and a highly centralised, CM-monitored bureaucracy. The effect "
        "has been to shift UP from Mandal-era caste bargaining toward a presidential-style, performance-and-Hindutva model "
        "that subordinates caste blocs \u2013 notably excluding Yadav patronage \u2013 within a larger Hindu coalition. It "
        "is durable, but vulnerable when jobs and farm distress dominate the agenda, as 2024 demonstrated.",
    ]),
    ("Current Challenges", [
        "Recommendation: UP's governance deficits \u2013 unemployment, agrarian distress, infrastructure and "
        "human-development gaps, and contested law-and-order claims \u2013 are now the opposition's main opening, and "
        "the party that credibly owns delivery on these issues gains.",
        "Joblessness is the sharpest. The 2024 cancellation of both the UP Police constable exam (around 48 lakh "
        "aspirants) and the UPPSC RO/ARO exam (about 10 lakh) over paper leaks \u2013 178 FIRs and roughly 396 arrests "
        "\u2013 crystallised youth anger and fed the BJP's 2024 losses. Agrarian distress is acute in Bundelkhand "
        "(drought, debt, farmer suicides) and statewide through the stray-cattle menace created by the cow-slaughter "
        "crackdown, unpaid sugarcane dues and MSP demands championed by the Bharatiya Kisan Union. Despite flagship "
        "expressways and airports, per-capita income remains near half the national average, with persistent gaps in "
        "health (exposed during COVID-19), education and urban services. On law and order, the BJP claims improvement "
        "through its encounter and \u201cbulldozer\u201d model, while critics cite custodial deaths, communal flashpoints "
        "and crimes against women and Dalits.",
        "Electorally, the BJP has answered with a \u201clabharthi\u201d (beneficiary) strategy \u2013 free foodgrain to "
        "about 15 crore people, PM Awas housing, Ujjwala, toilets and Kisan Samman Nidhi \u2013 alongside its "
        "law-and-order plank; the SP-led opposition has answered with jobs, a caste census, defence of reservation and "
        "agrarian relief. The recommendation follows directly: incumbents must convert capital spending into employment "
        "and human-development outcomes, while challengers must pair grievance with a concrete delivery alternative "
        "rather than identity alone.",
    ]),
    ("Electoral Shifts", [
        "Recommendation (the storyline's spine): Congress fell because its coalition fragmented; the SP and BSP rose "
        "and then hit caste ceilings; the BJP won by transcending those ceilings; and 2024 showed that social-justice "
        "arithmetic plus governance discontent can still beat it \u2013 the template for 2027.",
        "Congress's decline. From dominance in the 1950s\u201380s (its last majorities in 1980 and 1985), Congress "
        "collapsed after 1989 when the twin shocks of Mandal (OBC reservation) and Mandir (the Ram movement) split its "
        "umbrella of Brahmins, Dalits and Muslims: OBCs left for the SP, Dalits for the BSP, upper castes and "
        "Hindu-consolidation voters for the BJP, and Muslims gravitated to whichever party could defeat the BJP. With no "
        "social core, an ossified organisation and no state-level leadership, Congress has not formed a UP government "
        "since 1989 and won just 2 of 403 seats in 2022.",
        "The BSP-SP see-saw and the 2017 rupture. The BSP's 206 (2007) and the SP's 224 (2012) majorities were built by "
        "widening a caste core \u2013 the BSP's Dalit base plus Brahmins (\u201csarvajan\u201d), the SP's MY base plus "
        "anti-incumbency. Both were swept away in 2017 (BJP 312) when the BJP fused non-Jatav Dalits and non-Yadav OBCs "
        "with upper castes: each rival stayed near its ~22% ceiling (BSP 22.2%, SP 21.8%) while the BJP's Hindu coalition "
        "reached 39.7%.",
        "The BJP's winning formula, 2014\u20132022. Seventy-one of 80 seats (42.6%) in 2014, 312 in 2017, 62 of 80 "
        "(about 50%) in 2019 and 255 in 2022 rested on three legs: caste engineering (non-Yadav OBC, non-Jatav Dalit and "
        "upper caste, with allies Apna Dal and NISHAD); welfare delivery (the labharthi vote); and nationalism and "
        "Hindutva (Modi's leadership, Balakot in 2019, the Ram temple) that subsumed caste under Hindu unity.",
        "The alliances and the 2024 turn. The 2017 SP-Congress pact (54 seats) and the 2019 SP-BSP-RLD Mahagathbandhan "
        "(15 of 80 \u2013 BSP 10, SP 5 \u2013 against the BJP's 62) failed because arithmetic never became chemistry: the "
        "Yadav and Jatav votes did not fully transfer, non-Jatav Dalits and non-Yadav OBCs stayed with the BJP, and "
        "Modi's nationalism and welfare overrode caste sums. 2024 changed because the SP's PDA (Pichhda-Dalit-Alpsankhyak) "
        "broadened the base beyond MY through disciplined ticket distribution to non-Yadav OBCs, Dalits and most-backward "
        "castes; the \u201cConstitution and reservation in danger\u201d narrative (against the BJP's \u201c400-paar\u201d "
        "slogan) mobilised Dalits and OBCs; jobs, paper-leak and farm anger peaked; and the SP-Congress vote transferred "
        "efficiently as the BSP collapsed to 9.4%. The result was INDIA 43, BJP 33.",
        "Lessons for 2027: (1) keep and widen the PDA coalition and never rely on the MY core alone; (2) build "
        "booth-level organisation to convert vote share into seats \u2013 the 2019 failure was chemistry, not arithmetic; "
        "(3) lead with local delivery issues (jobs, inflation, agrarian relief, law-and-order accountability) and a "
        "credible alternative to welfare, not identity alone; (4) keep social justice (caste census, reservation) central "
        "while absorbing the BSP's drifting Dalit vote; and (5) lock in disciplined seat-sharing and candidate selection "
        "early. For the BJP, the counter is the mirror image: convert spending into jobs and re-secure non-Yadav OBC and "
        "non-Jatav Dalit loyalty, because its Hindu-coalition ceiling is breached precisely when governance fails.",
    ]),
]


def style_run(run, size=BODY_SIZE, bold=False):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), BODY_FONT)


def main():
    doc = docx.Document()

    # base style
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # page geometry: A4, 0.5" margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.header_distance = Inches(0.3)

    # header: right-aligned, Arial 9
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(HEADER_TEXT)
    style_run(hr, size=9)

    for heading, paras in CONTENT:
        if heading:
            hpar = doc.add_paragraph()
            hpar.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hpar.paragraph_format.space_before = Pt(2)
            hpar.paragraph_format.space_after = Pt(4)
            hpar.paragraph_format.line_spacing = 1.0
            hpar.paragraph_format.keep_with_next = True
            r = hpar.add_run(heading)
            style_run(r, size=BODY_SIZE, bold=True)
        for i, text in enumerate(paras):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = GAP_AFTER
            r = p.add_run(text)
            style_run(r)

    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    main()
